# @author: Alexander Chapman, Douglas Hale
# Used to collect text from docx files and then move them to txt for easier retrieval when training

# Import necessary libraries
from scrapeSingleArticle import scrapeSingleArticle
import string
from nltk.corpus import stopwords
import os
import certifi
import docx


os.environ['SSL_CERT_FILE'] = certifi.where()

# Possibly go and look at removing file path similarity with readyModelData
class CollectTrainingData:
    """
        Collects body text of all docx files and writes to a new folder to create txt files. Then the files are read and the text is put into an array. Eventually will be a database. Or no model should write to a database everytime it gets new data.
    
        Assumptions:
            1) The source folders will match those listed in the @sources tuple, following the same top-down directory order as the left-to-right order in @sources.

            2) The values in @classifers correspond to the values in @sources that share the same index.
            
        Methods:
            @__call__(string, string(), string())
                Constructor that will collect all articles, clean their text, and append a tuple containing the cleaned text and its classifier to a list that will be returned.    

            @_convertToTxt(string, boolean)
                Will collect the bodies of the docx files collected using the Nexus Uni database and write them to txt files to make finding, reading, and editing the articles easier for the user. 
            
            @_createTXTFile(string, string)
                Helper method that creates a txt file by using the old folder path except changing the original directory to CleanedArticles and combining that with the filename except with docx changed to txt.

            @_writeSourceArticles(docx.Document, string)
                Helper method to write the body of articles to a given txt file.

            @_writeNBCArticles(docx.Document, string)
                Helper method for NBC articles to scrape the article url and write the body to a given txt file.

            @_readyModelData(string, string(), string())
                Creates a list of tuples containing article bodies and their classifiers, which will be used by the model for training and testing.
    """
    

    def __call__(self, directory_path, sources, classifiers):
        """
            Constructor that will collect all articles, clean their text, and append a tuple containing the cleaned text and its classifier to a list that will be returned.

            Parameters:
                @directory_path :: string
                    The path to the directory that contains the source folders of docx files. 
                @sources :: string()
                    Tuple that contains the names of the news sources. This should be in the same order as the folder. 
                @classifiers :: string()
                    Tuple that contains the classifications of the sources. They should correspond to source at same index in sources.
        """

        print("hit")
        for source in sources:
            source_path = os.path.join(directory_path, source)
            if source == "NBC":
                self._convertToTxt(source_path, True)
            else:
                self._convertToTxt(source_path, False)

        return self._readyModelData(directory_path, sources, classifiers)


    def _convertToTxt(self, source_path, isNBC):
        """
            Will collect the bodies of the docx files collected using the Nexus Uni database and write them to txt files to make finding, reading, and editing the articles easier for the user. 

            Parameters:
                @source_path :: string
                    The path to the folder containing the news sources' articles.
                @isNBC :: boolean
                    Tells if the source is NBC which has a different docx format.
        
            Exceptions:
                @docx.opc.exceptions.PackageNotFoundError
                    Caught if the file being opened is not a docx file.
        """

        # Loops through all folders at the source path and skips hidden folders like .DSStore that can be automatically added by the system
        for folder in os.listdir(source_path):       
            if folder.startswith("."):
                continue

            # Loops through all files in folder path and stores the complete path as a variable
            folder_path = os.path.join(source_path, folder)
            for filename in os.listdir(folder_path):
                file_path = os.path.join(folder_path, filename)
                
                # Tries to create a new Document object by using the file_path. Catches exception and skips file if it does not work
                try:
                    doc = docx.Document(file_path)

                    # Returns the new txt path that will be written to, by use of a helper method.
                    new_directory_path = self._createTXTFile(folder_path, filename)

                    # Checks if the source is NBC, which has a different way of scraping the article. Calls helper methods to write the txt files
                    if isNBC:
                        self._writeNBCFiles(doc, new_directory_path)
                    else:
                        self._writeSourceFiles(doc, new_directory_path)
                except docx.opc.exceptions.PackageNotFoundError:
                    print(f"Skipping file: {filename} (Not a valid .docx file)")



    def _createTXTFile(self, folder_path, filename):
        """
            Helper method that creates a txt file by using the old folder path except changing the original directory to CleanedArticles and combining that with the filename except with docx changed to txt.

            Parameters:
                @folder_path :: string
                    The path to the folder that contains the original docx file.
                @filename :: string
                    The filename of the document. Will have it's tag changed to txt.

            Returns:
                @new_directory_path :: string
                    The complete path to the new txt file.
        """

        # Changes the parent folder in the path and then creates a new folder at that destination if there is none
        new_directory_path = folder_path.replace("ScrapedArticles","CleanedArticles") 
        os.makedirs(new_directory_path, exist_ok=True)  

        # Changes the docx tag to a txt and then creates the entire new file path, then returns it
        txtfile = filename.replace("docx", "txt")
        if txtfile == filename:
            raise ValueError("Is not a docx file.")
        
        new_directory_path = os.path.join(new_directory_path, filename.replace("docx", "txt")) 
        return new_directory_path
    


    def _writeSourceFiles(self, doc, new_directory_path):
        """
            Helper method to write the body of articles to a given txt file.
        
            Parameters:
                @doc :: docx.Document Object
                    The docx file from which the article body is being collected from.
                @new_directory_path :: string
                    The complete path to the new txt file.
        """

        inBody = False # Used to tell if we are looking at the body of an article
        for para in doc.paragraphs:  

            # If not in the body than it checks if the current line of text enters it, then the loop is moved forward
            if not inBody:
                if para.text == "Body":
                    inBody = True
                continue
                        
            # A check for the ending of the body
            if "Load-Date" in para.text:
                break

            # Writes the current paragraph to a txt file 
            with open(new_directory_path, "w", encoding='utf8') as f:
                f.write((para.text))
        
        # If there was no article in the document then the path is removed and a ValueError is raised
        if not inBody:
            try:
                os.remove(new_directory_path)
            except FileNotFoundError:
                pass
            raise ValueError("The article has no body.")
 

            
    def _writeNBCFiles(self, doc, new_directory_path):
        """
            Helper method for NBC articles to scrape the article url and write the body to a given txt file.

            Parameters:
                @doc :: docx.Document Object
                    The docx file from which the article body is being collected from.
                @new_directory_path :: string
                    The complete path to the new txt file.
        """

        url_found = False # Used to check if the url has been reached

        # Checks for the text, which means the next paragraph will be the url, and continues the loop
        for para in doc.paragraphs:
            if para.text == "Click to view full-text":
                url_found = True
                continue

            if url_found:
                # Scrapes the article body of the url and writes it to a txt file 
                try:
                    with open(new_directory_path, "w", encoding='utf8') as f:
                        body = scrapeSingleArticle(para.text)[3]
                        f.write(body)
                        break 
                except TypeError:
                    os.remove(new_directory_path)

        # If there was no url to an article in the document then the path is removed and a ValueError is raised
        if not url_found:
            try:
                os.remove(new_directory_path)
            except FileNotFoundError:
                pass
            raise ValueError("There was no article url.")
  
        
    def _readyModelData(self, directory_path, sources, classifier):
        """
            Collects the body of all txt files and adds them to the data list, along with their classifier, which will be returned to the NaiveBayesModel.

            Parameters:
                @directory_path :: string
                    The path to the directory that contains the source folders of txt files. 
                @sources :: string()
                    Tuple that contains the names of the news sources. This should be in the same order as the folder. 
                @classifiers :: string()
                    Tuple that contains the classifications of the sources. They should correspond to source at same index in sources.

            Returns:
                @x_test :: string[]
                    A list of article bodies that make up the first index of the returned tuple.
                @y_test :: string[]
                    A list of the corresponding classifiers for the x_test bodies. This makes up the second index of the returned tuple.
        """

        x_test = [] 
        y_test = []
        numFolder = 0 # Used to know the classifier for the current source folder

        # Loops through every source folder
        for source in sources:
            source_path = os.path.join(directory_path, source)

            # Loops through all folders at the source path and skips hidden folders like .DSStore that can be automatically added by the system
            for folder in os.listdir(source_path):       
                if folder.startswith("."):
                    continue

                # Loops through all files in folder path and stores the complete path as a variable
                folder_path = os.path.join(source_path, folder)
                for filename in os.listdir(folder_path):
                    file_path = os.path.join(folder_path, filename)

                    # Open the file and read the text to body
                    with open(file_path, 'r', encoding='utf8') as f:
                        body = f.read()

                    # Append the body to x_test and its corresponding classifier to y_test
                    x_test.append(body)
                    y_test.append(classifier[numFolder])

            # Add 1 to numFolder after the source folder is looped through completely
            numFolder += 1
        return x_test, y_test



    def _cleanText(self, text):
        """
            Cleans text by removing punctuation and stop words, before lowercasing the text.
            
            Parameters:
                @text :: string
                    The text that is going to be cleaned.

            Returns:
                @cleanedText :: string
                    The newly cleaned text.
        """
        text = text.translate(str.maketrans('', '', string.punctuation))
        text = text.lower()
        words = text.split()
        words = [w for w in words if w not in stopwords.words('english')]
        cleanedText = ' '.join(words)  # Join words back into a string
        return cleanedText

