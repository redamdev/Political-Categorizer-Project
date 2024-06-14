# @author: Alexander Chapman, Douglas Hale
# Will take the data and use it to train the model and complete testing

# Import necessary libraries
from scrapeSingleArticle import scrapeSingleArticle
import os
import re
import certifi
import docx
from sklearn.feature_extraction.text import CountVectorizer, TfidfVectorizer
from sklearn.naive_bayes import MultinomialNB
from sklearn.metrics import accuracy_score, classification_report
from sklearn.model_selection import train_test_split
from sklearn.metrics import confusion_matrix
from joblib import dump, load
os.environ['SSL_CERT_FILE'] = certifi.where()


""" Every time a user enters an input a model will be loaded and a question asked"""
class TrainingNBModel:
    """
        A Naive Bayes model where it can be created from scratch with given data or loaded into a saved model.
    
        Class Variables:
            @vectorizer :: CountVectorizer Object
                Used to transform data into vectors of 2 and 3 word phrases and phrase counts, which will be used by the model when training and testing.
            @model :: MultinomialNB Object
                Will be trained using a multinomial naive bayes algorithm and used to make predictions.
        
        Instance Variables:
            @x_test :: string()
                A tuple of the articles taken from the initial dataset and chosen for testing.
            @y_test :: string()
                A tuple of the classifiers taken from the initial dataset and chosen for testing.

        Methods:
            @__init__()
                Constructor that will declare and initialize the instance variables, @x_test and @y_test, to empty tuples.

            @trainModel(string, string(), string())
                Will use a given directory path with news sources and classifiers to collect the data and then split it into training and testing.

            @load_model()
                Load the model and vectorizer by using the predetermined paths

            @testTrainingAccuracy()
                Uses the instance variables x_test and y_test to test model accuracy while also reporting its Confusion Matrix and Classification Report.

            @_convertToTxt(string, boolean)
                Will collect the bodies of the docx files collected using the Nexus Uni database and write them to txt files to make finding, reading, and editing the articles easier for the user. 

            @_writeSourceArticles(docx.Document, string)
                Helper method to write the body of articles to a given txt file.

            @_writeNBCArticles(docx.Document, string)
                Helper method for NBC articles to scrape the article url and write the body to a given txt file.

            @_readyModelData(string, string(), string())
                Creates a list of tuples containing article bodies and their classifiers, which will be used by the model for training and testing.
    """

    # Class Variables
    #vectorizer = TfidfVectorizer(ngram_range=(1,1))
    vectorizer = CountVectorizer(ngram_range=(4,4))
    
    model = MultinomialNB()

    
    def __init__(self):
        """
            Constructor that will declare and initialize the instance variables, @x_test and @y_test, to empty tuples.

            Instance Variables:
                @x_test :: string()
                    A tuple of the articles taken from the initial dataset and chosen for testing.
                @y_test :: string()
                    A tuple of the classifiers taken from the initial dataset and chosen for testing.
        """

        self.x_test = () 
        self.y_test = ()


    def trainOldDataModel(self, sources):
        trainingData = _readyModelData("NewsBuster v2/backend/Model/Articles/CleanedArticles/", sources)

        x, y = trainingData

        # Seperates data into being for training or test with 80% of articles being used for training.
        x_train, self.x_test, y_train, self.y_test = train_test_split(x, y, test_size=0.2, random_state=42)

        # Transform training articles into vector form and have the vectorizer remember these by adding 'fit'.
        vectorized_x_train = self.vectorizer.fit_transform(x_train)

        # Trains the model with the given articles and classifiers, then saves and returns it.
        self.model.fit(vectorized_x_train, y_train)
        dump(self.model, 'NewsBuster v2/backend/Model/nb_model.joblib')
        dump(self.vectorizer, 'NewsBuster v2/backend/Model/vectorizer.joblib')


    def trainNewDataModel(self, directory_path, sources):
        """
            Creates and train a new multinomial naive bayes model with news articles by given sources and their given classifiers. The path to the articles is then used to collect training data from every source.

            Parameters:
                @directory_path :: string
                    The path to the directory that contains the source folders of docx files. 
                @sources :: string()
                    Tuple that contains the names of the news sources. This should be in the same order as the folder. 
                @classifiers :: string()
                    Tuple that contains the classifications of the sources. They should correspond to source at same index in sources.
        """
        
        # Collects the data and returns it in a 2D list with [article body, classifier] being the inside form.

        for source in sources:
            source_path = os.path.join(directory_path, source)
            if source == "NBC":
                _convertToTxt(source_path, True)
            else:
                _convertToTxt(source_path, False)

        self.trainOldDataModel(sources)


    def load_model(self):
        """
            Takes a file path as the argument and uses it to load that model into the current session.
    
            Parameters:
                @file_path :: string
                    The path to the file that stores the model.

            Returns:
                @self.model :: MultibinomialNB Object
                    The model that has been loaded.
            
        """
        self.model = load('NewsBuster v2/backend/Model/nb_model.joblib')
        self.vectorizer = load("NewsBuster v2/backend/Model/vectorizer.joblib") 
    

    def testTrainingAccuracy(self):
        """
            Uses the instance variables x_test and y_test to test model accuracy while also reporting its Confusion Matrix and Classification Report. 
        """

        # Vectorizes the x_test data for use by the model, before it is then used by the model to predict the classifiers.
        vectorized_x_test = self.vectorizer.transform(self.x_test)
        y_pred = self.model.predict(vectorized_x_test) 

        # Collects how accurate the model was by comparing the predicted classifiers against the actual ones. 
        accuracy = accuracy_score(self.y_test, y_pred)

        # Shows the exact numbers of predictions made for each classifier compared to actual results. 
        conf_matrix = confusion_matrix(self.y_test, y_pred) 

        # Can tell how precise the model was and how often it recalled the classes.
        class_report = classification_report(self.y_test, y_pred)

        # Prints results
        print(f"Model accuracy on test data: {accuracy:.2f}") 
        print("Confusion Matrix:\n", conf_matrix)
        print("Classification Report:\n", class_report)



def _convertToTxt(source_path, isNBC):
    """
        Will collect the bodies of the docx files collected using the Nexus Uni database and write them to txt files to make finding, reading, and editing the articles easier for the user. 

        Parameters:
            @source_path :: string
                The path to the folder containing the news sources' articles.
            @isNBC :: boolean
                Tells if the source is NBC which has a different docx format.
    
        Exceptions:
            @docx.opc.exceptions.PackageNotFoundError
                Caught if the file being opened is not a docx file.
    """

    # Loops through all folders at the source path and skips hidden folders like .DSStore that can be automatically added by the system
    for folder in os.listdir(source_path):       
        if folder.startswith("."):
            continue

        # Loops through all files in folder path and stores the complete path as a variable
        folder_path = os.path.join(source_path, folder)
        for filename in os.listdir(folder_path):
            file_path = os.path.join(folder_path, filename)

            # Finds the file that lists all of the articles in the folder and then deletes it before moving on
            if(filename == "Files (100)_doclist.docx"):
                print(file_path)
                os.remove(file_path)
                continue

            # Tries to create a new Document object by using the file_path. Catches exception and skips file if it does not work
            try:
                doc = docx.Document(file_path)

                # Changes the parent folder in the path and then creates a new folder at that destination if there is none
                new_directory_path = folder_path.replace("ScrapedArticles","CleanedArticles") 
                os.makedirs(new_directory_path, exist_ok=True)  

                # Changes the docx tag to a txt and then creates the entire new file path, then returns it
                txtfile = filename.replace("docx", "txt")
                if txtfile == filename:
                    raise ValueError("Is not a docx file.")
    
                new_directory_path = os.path.join(new_directory_path, filename.replace("docx", "txt")) 

                # Checks if the source is NBC, which has a different way of scraping the article. Calls helper methods to write the txt files
                if isNBC:
                    writeNBCFiles(doc, new_directory_path)
                else:
                    writeSourceFiles(doc, new_directory_path)
            except docx.opc.exceptions.PackageNotFoundError:
                print(f"Skipping file: {filename} (Not a valid .docx file)")



def writeSourceFiles(doc, new_directory_path):
    """
        Helper method to write the body of articles to a given txt file.
    
        Parameters:
            @doc :: docx.Document Object
                The docx file from which the article body is being collected from.
            @new_directory_path :: string
                The complete path to the new txt file.
    """
    if "/AP/" in new_directory_path:
        inBody = True
    else:
        inBody = False # Used to tell if we are looking at the body of an article
    body = ""
    for para in doc.paragraphs:  

        # If not in the body than it checks if the current line of text enters it, then the loop is moved forward
        if not inBody:
            if para.text == "Body":
                inBody = True
            continue
                    
        # A check for the ending of the body
        if "Load-Date" in para.text:
            break

        body += para.text

    # Writes the current paragraph to a txt file 
    with open(new_directory_path, "w", encoding='utf-8') as f:
        f.write((body))
    
    # If there was no article in the document then the path is removed and a ValueError is raised
    if not inBody:
        try:
            os.remove(new_directory_path)
        except FileNotFoundError:
            pass
        raise ValueError("The article has no body.")


        
def writeNBCFiles(doc, new_directory_path):
    """
        Helper method for NBC articles to scrape the article url and write the body to a given txt file.

        Parameters:
            @doc :: docx.Document Object
                The docx file from which the article body is being collected from.
            @new_directory_path :: string
                The complete path to the new txt file.
    """

    url_found = False # Used to check if the url has been reached

    # Checks for the text, which means the next paragraph will be the url, and continues the loop
    for para in doc.paragraphs:
        if para.text == "Click to view full-text":
            url_found = True
            continue

        if url_found:
            # Scrapes the article body of the url and writes it to a txt file 
            try:
                with open(new_directory_path, "w", encoding='utf-8') as f:
                    body = scrapeSingleArticle(para.text)[3]
                    f.write(body)
                    break 
            except TypeError:
                os.remove(new_directory_path)

    # If there was no url to an article in the document then the path is removed and a ValueError is raised
    if not url_found:
        try:
            os.remove(new_directory_path)
        except FileNotFoundError:
            pass
        raise ValueError("There was no article url.")

    
def _readyModelData(directory_path, sources):
    """
        Collects the body of all txt files and adds them to the data list, along with their classifier, which will be returned to the NaiveBayesModel.

        Parameters:
            @directory_path :: string
                The path to the directory that contains the source folders of txt files. 
            @sources :: string()
                Tuple that contains the names of the news sources. This should be in the same order as the folder. 
            @classifiers :: string()
                Tuple that contains the classifications of the sources. They should correspond to source at same index in sources.

        Returns:
            @x_test :: string[]
                A list of article bodies that make up the first index of the returned tuple.
            @y_test :: string[]
                A list of the corresponding classifiers for the x_test bodies. This makes up the second index of the returned tuple.
    """

    x_test = [] 
    y_test = []

    # Loops through every source folder
    for source in sources:
        if "Manual" in source:
            bodies, classifiers = readyManualData(source)
            x_test.extend(bodies)
            y_test.extend(classifiers)
            continue

        source_path = os.path.join(directory_path, source)

        # Loops through all folders at the source path and skips hidden folders like .DSStore that can be automatically added by the system
        for folder in os.listdir(source_path):       
            if folder.startswith("."):
                continue

            # Loops through all files in folder path and stores the complete path as a variable
            folder_path = os.path.join(source_path, folder)
            for filename in os.listdir(folder_path):
                file_path = os.path.join(folder_path, filename)

                # Open the file and read the text to body then replace all periods with spaces
                with open(file_path, 'r+', encoding='utf-8') as f:
                    body = f.read() 
                    f.seek(0) # Moves file ptr back to start
                    body = body.replace(".", " ")
                    f.write(body)
                    f.truncate() # Erases any extra bytes if new read is smaller than previous
            
                # Append the body to x_test and its corresponding classifier to y_test
                x_test.append(body)
                y_test.append(sources[source])
    return x_test, y_test



def addManualChecks(url, classifier):
    directoryPath = "NewsBuster v2/backend/Model/Articles/CleanedArticles/"
    article = scrapeSingleArticle(url)
    title = article[0]
    body = article[3]

    print(body)
    if classifier == "right":
        folder_path = os.path.join(directoryPath, "RightManual")
        file = title + ".txt"
        file_path = os.path.join(folder_path, file)
        with open(file_path, "w", encoding='utf-8') as f:
            f.write((body)) 

    elif classifier == "left":
        folder_path = os.path.join(directoryPath, "LeftManual")
        file = title + ".txt"
        file_path = os.path.join(folder_path, file)
        with open(file_path, "w", encoding='utf-8') as f:
            f.write((body))  

    elif classifier == "middle":
        folder_path = os.path.join(directoryPath, "MiddleManual")
        file = title + ".txt"
        file_path = os.path.join(folder_path, file)
        with open(file_path, "w", encoding='utf-8') as f:
            f.write((body))  

def readyManualData(source):
    directoryPath = "NewsBuster v2/backend/Model/Articles/CleanedArticles/"
    folder_path = os.path.join(directoryPath, source)
 
    x_test = [] 
    y_test = []

    if "Right" in source:
        classifier = "right"
    elif "Left" in source:
        classifier = "left"
    elif "Middle" in source:
        classifier = "middle"

    for filename in os.listdir(folder_path):
        if filename.startswith("."):
            continue
        file_path = os.path.join(folder_path, filename)

        # Open the file and read the text to body then replace all periods with spaces
        with open(file_path, 'r+', encoding='utf-8') as f:
            body = f.read() 
            f.seek(0) # Moves file ptr back to start
            body = body.replace(".", " ")
            f.write(body)
            f.truncate() # Erases any extra bytes if new read is smaller than previous

        x_test.append(body)
        y_test.append(classifier)

    return x_test, y_test

def delete_files_with_multiples(directory, pattern_to_delete):
    for source in os.listdir(directory):
        source_path = os.path.join(directory, source)
        if "Manual" in source:
             continue
        if source.startswith("."):
            continue 
        # Loops through all folders at the source path and skips hidden folders like .DSStore that can be automatically added by the system
        for folder in os.listdir(source_path):       
            if folder.startswith("."):
                continue
            
            # Loops through all files in folder path and stores the complete path as a variable
            folder_path = os.path.join(source_path, folder)
            for filename in os.listdir(folder_path):
                file_path = os.path.join(folder_path, filename)
    
                if re.search(pattern_to_delete, filename):
                    try:
                        os.remove(file_path)
                        print(f"Deleted: {file_path}")
                    except OSError as e:
                        print(f"Error deleting {file_path}: {e}")

# Usage Example
directory_to_clean = "NewsBuster v2/backend/Model/Articles/CleanedArticles/"
pattern = r"\(\d\)"  # Regular expression to match any single digit within parentheses



newsSources = {
                "AP" : "middle",
                "CNN" : "left",
                "NBC" : "left", 
                "Forbes" : "middle", 
                "Newsweek" : "right", 
                "Daily_Caller" : "right",
                "RightManual" : "right",
                "LeftManual" : "left",
                "MiddleManual" : "middle"
            }


#addManualChecks("https://www.cnn.com/2024/03/17/politics/dark-money-fga-ashcroft-invs/index.html", "left")

model = TrainingNBModel()
model.trainOldDataModel(newsSources)
#model.trainNewDataModel("NewsBuster v2/backend/Model/Articles/ScrapedArticles/", newsSources)
#model.load_model()
model.testTrainingAccuracy()